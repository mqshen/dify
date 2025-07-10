
from typing import IO
from docx import Document

def find_lcs(str1, str2):
    lstr1 = len(str1)
    lstr2 = len(str2)
    record = [[0 for i in range(lstr2 + 1)] for j in range(lstr1 + 1)]  # 多一位
    maxNum = 0
    p = 0
    for i in range(lstr1):
        for j in range(lstr2):
            if str1[i] == str2[j]:
                record[i + 1][j + 1] = record[i][j] + 1
                if record[i + 1][j + 1] > maxNum:
                    maxNum = record[i + 1][j + 1]
                    p = i + 1

    return str1[p - maxNum: p], maxNum

class DocxTemplateRender(object):
    def __init__(self, filepath: str = None, file_content: IO[bytes] = None):
        self.filepath = filepath
        self.file_content = file_content
        if self.filepath:
            self.doc = Document(self.filepath)
        else:
            self.doc = Document(self.file_content)

    def render(self, template_def: dict[str, any]):
        doc = self.doc

        # 采用找key 方式
        for k1, v1 in template_def.items():
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if k1 in cell.text:
                            for one in cell.paragraphs:
                                if k1 in one.text:
                                    new_text = one.text.replace("{{" + k1 + "}}", v1)
                                    one.runs[0].text = new_text
                                    for r_index, r in enumerate(one.runs):
                                        if r_index == 0:
                                            continue
                                        r.text = ''

            for p in doc.paragraphs:
                # p.text = p.text.replace(k1, v1)
                # continue
                if k1 not in p.text:
                    continue
                runs_cnt = len(p.runs)
                s_e = []
                i = 0
                while i < runs_cnt:
                    new_i = i + 1
                    for j in range(i + 1, runs_cnt + 1):
                        part_text = ''.join([r.text for r in p.runs[i:j]])
                        if k1 in part_text:
                            # 找到最小的范围内包含k1的runs
                            tmp_i, tmp_j = i, j
                            while tmp_i <= tmp_j:
                                tmp_part_text = ''.join([r.text for r in p.runs[tmp_i:tmp_j]])
                                if k1 in tmp_part_text:
                                    tmp_i += 1
                                    continue
                                else:
                                    tmp_i -= 1
                                    break
                            s_e.append((tmp_i, j))
                            new_i = j
                            break
                    i = new_i

                for one in s_e:
                    s, e = one
                    assert e > 0, [r.text for r in p.runs]
                    # tgt_text = [r.text for r in p.runs[s:e]]
                    if e - s == 1:
                        replace_mapping = [(k1, v1)]
                    elif e - s == 2:
                        s_tgt_text = p.runs[s].text
                        comm_str, max_num = find_lcs(k1, s_tgt_text)
                        assert k1.startswith(comm_str)
                        p1 = comm_str
                        p2 = k1[max_num:]
                        n = len(v1)
                        sub_n1 = int(1.0 * p1 / (p1 + p2) * n)
                        # sub_n2 = n - sub_n1
                        replace_mapping = [(p1, v1[:sub_n1]), (p2, v1[sub_n1:])]
                    elif e - s == 3:
                        m_text = p.runs[s + 1].text
                        head_tail = k1.split(m_text, 1)
                        assert len(head_tail) == 2
                        h_text = head_tail[0]
                        t_text = head_tail[1]
                        replace_mapping = [(h_text, ''), (m_text, v1), (t_text, '')]
                    else:
                        m_texts = [p.runs[i].text for i in range(s + 1, e - 1)]
                        m_text = ''.join(m_texts)
                        head_tail = k1.split(m_text, 1)
                        assert len(head_tail) == 2
                        h_text = head_tail[0]
                        t_text = head_tail[1]
                        replace_mapping = [(h_text, '')]
                        replace_mapping.append((m_texts[0], v1))
                        for text in m_texts[1:]:
                            replace_mapping.append((text, ''))

                        replace_mapping.append((t_text, ''))

                    for i in range(s, e):
                        _k, _v = replace_mapping[i - s]
                        p.runs[i].text = p.runs[i].text.replace("{{" + _k + "}}", _v)


        return doc
